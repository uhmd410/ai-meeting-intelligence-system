"""
export_service.py — Generate PDF and DOCX exports of meeting minutes.

Uses fpdf2 for PDF generation and python-docx for DOCX generation.
Both functions accept the dict returned by crud.get_meeting_detail()
and return raw bytes suitable for a StreamingResponse.
"""

from io import BytesIO
from datetime import datetime

from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _format_generated_at(dt) -> str:
    """Format a datetime (or ISO string) as 'Month DD, YYYY at HH:MM'."""
    if dt is None:
        return "N/A"
    if isinstance(dt, str):
        # Handle ISO format strings from JSON serialization
        dt = datetime.fromisoformat(dt.replace("Z", "+00:00"))
    return dt.strftime("%B %d, %Y at %H:%M")


# ─── PDF Export ────────────────────────────────────────────────────────────────


def export_to_pdf(meeting_detail: dict) -> bytes:
    """Generate a professional PDF of the meeting minutes and return raw bytes."""
    title = meeting_detail["title"]
    minutes = meeting_detail.get("minutes")
    action_items = meeting_detail.get("action_items", [])
    generated_at = minutes["generated_at"] if minutes else None

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # ── Header ──────────────────────────────────────────────────
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(
        0,
        6,
        f"Generated: {_format_generated_at(generated_at)}",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.set_text_color(0, 0, 0)
    pdf.ln(6)

    # Helper: section heading
    def _section_heading(text: str):
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_fill_color(240, 240, 245)
        pdf.cell(0, 9, f"  {text}", new_x="LMARGIN", new_y="NEXT", fill=True)
        pdf.ln(3)

    # Helper: bullet list
    def _bullet_list(items: list | None):
        pdf.set_font("Helvetica", "", 11)
        if not items:
            pdf.cell(0, 7, "    None recorded.", new_x="LMARGIN", new_y="NEXT")
        else:
            for item in items:
                x = pdf.get_x()
                pdf.cell(6, 7, "-")  # bullet marker
                # multi_cell for wrapping long text; effective_width accounts for indent
                effective_width = pdf.w - pdf.l_margin - pdf.r_margin - 6
                pdf.multi_cell(effective_width, 7, f" {item}")
        pdf.ln(2)

    # ── Summary ─────────────────────────────────────────────────
    _section_heading("Summary")
    pdf.set_font("Helvetica", "", 11)
    summary_text = minutes["summary"] if minutes and minutes.get("summary") else "None recorded."
    pdf.multi_cell(0, 7, summary_text)
    pdf.ln(2)

    # ── Discussion Points ───────────────────────────────────────
    _section_heading("Discussion Points")
    _bullet_list(minutes.get("discussion_points") if minutes else None)

    # ── Decisions ───────────────────────────────────────────────
    _section_heading("Decisions")
    _bullet_list(minutes.get("decisions") if minutes else None)

    # ── Action Items (table) ────────────────────────────────────
    _section_heading("Action Items")
    if not action_items:
        pdf.set_font("Helvetica", "", 11)
        pdf.cell(0, 7, "    None recorded.", new_x="LMARGIN", new_y="NEXT")
    else:
        col_widths = [90, 40, 40]  # Task, Owner, Due Date
        table_width = sum(col_widths)
        headers = ["Task", "Owner", "Due Date"]

        # Header row
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_fill_color(55, 65, 81)
        pdf.set_text_color(255, 255, 255)
        for i, header in enumerate(headers):
            pdf.cell(col_widths[i], 8, f" {header}", border=1, fill=True)
        pdf.ln()

        # Data rows
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(0, 0, 0)
        for idx, ai in enumerate(action_items):
            if idx % 2 == 0:
                pdf.set_fill_color(249, 250, 251)
            else:
                pdf.set_fill_color(255, 255, 255)
            pdf.cell(col_widths[0], 7, f" {ai.get('task', '')}"[:55], border=1, fill=True)
            pdf.cell(col_widths[1], 7, f" {ai.get('owner', '')}", border=1, fill=True)
            pdf.cell(col_widths[2], 7, f" {ai.get('due_date', '')}", border=1, fill=True)
            pdf.ln()
    pdf.ln(2)

    # ── Risks ───────────────────────────────────────────────────
    _section_heading("Risks")
    _bullet_list(minutes.get("risks") if minutes else None)

    # ── Next Steps ──────────────────────────────────────────────
    _section_heading("Next Steps")
    _bullet_list(minutes.get("next_steps") if minutes else None)

    return bytes(pdf.output())


# ─── DOCX Export ───────────────────────────────────────────────────────────────


def export_to_docx(meeting_detail: dict) -> bytes:
    """Generate a professional DOCX of the meeting minutes and return raw bytes."""
    title = meeting_detail["title"]
    minutes = meeting_detail.get("minutes")
    action_items = meeting_detail.get("action_items", [])
    generated_at = minutes["generated_at"] if minutes else None

    doc = Document()

    # ── Adjust default style ────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # ── Header ──────────────────────────────────────────────────
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(
        f"Generated: {_format_generated_at(generated_at)}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(120, 120, 120)

    # Helper: add a section heading
    def _add_section(text: str):
        doc.add_heading(text, level=2)

    # Helper: add a bullet list, or "None recorded."
    def _add_bullets(items: list | None):
        if not items:
            p = doc.add_paragraph("None recorded.")
            p.style = doc.styles["Normal"]
            p.runs[0].font.italic = True
        else:
            for item in items:
                doc.add_paragraph(item, style="List Bullet")

    # ── Summary ─────────────────────────────────────────────────
    _add_section("Summary")
    summary_text = minutes["summary"] if minutes and minutes.get("summary") else "None recorded."
    p = doc.add_paragraph(summary_text)
    if summary_text == "None recorded.":
        p.runs[0].font.italic = True

    # ── Discussion Points ───────────────────────────────────────
    _add_section("Discussion Points")
    _add_bullets(minutes.get("discussion_points") if minutes else None)

    # ── Decisions ───────────────────────────────────────────────
    _add_section("Decisions")
    _add_bullets(minutes.get("decisions") if minutes else None)

    # ── Action Items (table) ────────────────────────────────────
    _add_section("Action Items")
    if not action_items:
        p = doc.add_paragraph("None recorded.")
        p.runs[0].font.italic = True
    else:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True

        # Header row
        hdr_cells = table.rows[0].cells
        for i, label in enumerate(["Task", "Owner", "Due Date"]):
            hdr_cells[i].text = label
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True

        # Data rows
        for ai in action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = ai.get("task", "")
            row_cells[1].text = ai.get("owner", "")
            row_cells[2].text = ai.get("due_date", "")

    # ── Risks ───────────────────────────────────────────────────
    _add_section("Risks")
    _add_bullets(minutes.get("risks") if minutes else None)

    # ── Next Steps ──────────────────────────────────────────────
    _add_section("Next Steps")
    _add_bullets(minutes.get("next_steps") if minutes else None)

    # ── Serialize to bytes ──────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
